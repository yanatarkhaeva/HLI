import re
from docx import Document
from io import BytesIO


class KnowledgeBase():
    def __init__(self, working_memory):
        self.working_memory = working_memory
        with open("Produktsionnye_pravila.docx", 'rb') as file:
            source_stream = BytesIO(file.read())
        self.doc = Document(source_stream)
        self.rules = []
        self.parse_facts()
        self.parse_rules()

        pass

    def add_rule(self, name, preconditions, insert):
        rule = KnowledgeBase.Rule(name, preconditions, insert)
        self.rules.append(rule)

    def add_question(self, name, preconditions, question):
        question = KnowledgeBase.Question(name, preconditions, question)
        self.rules.append(question)

    def get_rules(self):
        return self.rules

    def parse_facts(self):
        for paragraph in self.doc.paragraphs:
            if 'IF «' in paragraph.text:
                new_str = str(paragraph.text)
                res = re.split(r'[«»]', new_str)
                # print(res)
                for i, split in enumerate(res, 0):
                    if "IF" in split:
                        self.working_memory.add_fact(res[i + 1])
                    elif " AND " in split:
                        self.working_memory.add_fact(res[i + 1])
                    elif " THEN " in split:
                        self.working_memory.add_fact(res[i + 1])
                pass

    def parse_rules(self):
        name = None
        preconditions = {}
        insert = {}
        fact_key = ""
        insert_key = ""
        question = ""
        flag_then = False
        flag_ask = False
        for paragraph in self.doc.paragraphs:
            new_str = str(paragraph.text)
            res = re.split(r'[«»\n]', new_str)
            # print(res)
            for i, split in enumerate(res, 0):
                if "Правило" in split:
                    name = split[9:]
                elif "Вопрос" in split:
                    name = split[8:]
                elif ("IF" in split) or ("AND" in split):
                    fact_key = res[i + 1]
                    # print("key: " + fact_key)
                    pass
                elif ("=" in split) and (flag_then is False):
                    fact_value = res[i + 1]
                    preconditions[fact_key] = fact_value
                    # print("value: " + fact_value + "\n---")
                elif "THEN" in split:
                    insert_key = res[i + 1]
                    # print("insert key: " + insert_key + "\t")
                    flag_then = True
                elif ("=" in split) and (flag_then is True):
                    insert_value = res[i + 1]
                    insert[insert_key] = insert_value
                    # print("insert value: " + insert_value + "\n00000000000000000000000000")
                    flag_then = False
                    pass
                elif "ASK" in split:
                    question = res[i + 1]
                    flag_ask = True
                    # print("question: " + question)
                pass
            if name is not None and flag_ask is False:
                self.add_rule(name, preconditions, insert)
            elif name is not None and flag_ask is True:
                self.add_question(name, preconditions, question)
            flag_ask = False
            name = None
            preconditions.clear()
            insert.clear()
            question = None

        pass

    class Rule:
        def __init__(self, name, preconditions, insert):
            self.name = name
            self.preconditions = {}
            self.preconditions.update(preconditions)
            self.is_used = False
            self.insert = {}
            self.question = "Так как " + str(str(self.preconditions)[2:-2].translate({ord('\''): ord(' '), ord(':'): ord('=')}) + ", следовательно " + str(insert)[2:-2].translate({ord('\''): ord(' '), ord(':'): ord('=')}))
            try:
                self.insert.update(insert)
            except:
                self.insert.update("")

        def update_facts(self, facts, Answer=None):
            # обновление словаря фактов выводом который сформировался благодаря действию правила
            facts.update(self.insert)
            pass

    class Question(Rule):
        def __init__(self, name, preconditions, question):
            insert={}
            insert[name] = None
            super().__init__(name, preconditions, insert)
            self.question = question
            pass

        def update_facts(self, facts, Answer=None):
            answer = Answer
            if answer == "да" or answer == "нет":
                facts[self.name] = answer
            elif "ЧСС" in self.name:
                if 60 > int(answer):
                    facts["ЧСС"] = "<60"
                elif 90 < int(answer):
                    facts["ЧСС"] = ">90"
                elif 60 < int(answer) < 90:
                    facts["ЧСС"] = "норма"
            pass

        def print_question(self):
            return self.question

        pass


class IKnowledgeBase(KnowledgeBase):
    # working_memory = zi.Attribute("""Ссылка на рабочую память""")
    # doc = zi.Attribute("""Загруженный документ для парсера""")

    def __init__(self, working_memory):
        super().__init__(working_memory)

    def parse_facts(self):
        """Запоминаем только факты"""
        super().parse_facts()
        pass

    def parse_rules(self):
        """Запоминаем правила"""
        super().parse_rules()

    pass